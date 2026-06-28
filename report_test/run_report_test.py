"""Headless report test harness — exercises the real desktop report pipeline.

Drops the UI entirely and calls the SAME engine + report builder the PyWebview
shell uses (``core.analysis`` → ``desktop.report_host.build_report``), so what
this produces is byte-for-byte what the app would hand to its Save dialog.
Because it is pure Python it does NOT touch ``web/dist/`` — it can never be
fooled by a stale compiled bundle.

Jobs model
----------
Define one or more report runs in ``report_test/jobs.json``. Each job picks a
logger file + template from ``report_test/fixtures/`` and (optionally) a time
window to slice a single CSV into one test. This mirrors the app's
``time_start`` / ``time_end`` filtering, so a CSV holding several back-to-back
tests can be split into a report per test.

    {
      "jobs": [
        {
          "name": "block_load",
          "csv": "AWS54-1.4B-4556 G3.csv",
          "template": "4.01C ... BLOCK LOAD TEMPLATE.docx",
          "time_start": "2026-05-25 08:51:33",
          "time_end":   "2026-05-25 08:55:19",
          "fields": {"report_title": "...", "gen_sn": "..."},
          "config": {"nominal_voltage": 415.0}
        }
      ]
    }

Outputs land in ``report_test/output/<name>/``:
    report.pdf       — HTML→PDF pipeline (Chrome headless)
    report.docx      — Word output (only with a template)
    report.docx.pdf  — LibreOffice render of the .docx, for visual inspection
    report.html      — self-contained HTML
    summary.json     — analysis contract + warnings + image-placement audit

Usage
-----
    .venv/bin/python report_test/run_report_test.py            # run all jobs
    .venv/bin/python report_test/run_report_test.py --job iso  # one job by name
    .venv/bin/python report_test/run_report_test.py --watch    # re-run on change

Single-shot CLI mode (no jobs.json needed) is still available:
    ... --csv FILE --template FILE --time-start "..." --time-end "..." --name foo
"""
from __future__ import annotations

import argparse
import base64
import io
import json
import os
import subprocess
import sys
import time

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.dirname(_THIS_DIR)
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

FIXTURES_DIR = os.path.join(_THIS_DIR, "fixtures")
OUTPUT_DIR = os.path.join(_THIS_DIR, "output")
JOBS_FILE = os.path.join(_THIS_DIR, "jobs.json")

import shutil as _shutil

# NOTE: the harness deliberately does NOT set PQA_CHROMIUM. It relies on
# report_host.find_chromium() — the same browser discovery the desktop app uses —
# so the harness exercises the real PDF path and would catch a discovery gap
# (e.g. macOS .app bundles not being found) instead of masking it with an env var.


def _find_soffice() -> str | None:
    """Locate LibreOffice ``soffice`` for the .docx→PDF visual-render step.

    Cross-platform: PATH first, then the standard install location per OS.
    Returns None if not installed (the harness then skips report.docx.pdf and
    still produces report.pdf via the HTML/Chromium pipeline).
    """
    found = _shutil.which("soffice") or _shutil.which("soffice.exe")
    if found:
        return found
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",        # Windows
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",     # macOS
        "/usr/bin/soffice", "/usr/bin/libreoffice",                 # Linux
    ):
        if os.path.exists(p):
            return p
    return None


_SOFFICE = _find_soffice()

# Source files whose changes should retrigger --watch (the report pipeline).
_WATCH_SOURCES = [
    os.path.join(_REPO_ROOT, "core"),
    os.path.join(_REPO_ROOT, "desktop"),
    os.path.join(_REPO_ROOT, "report.py"),
    os.path.join(_REPO_ROOT, "html_report.py"),
    os.path.join(_REPO_ROOT, "visualizations.py"),
    FIXTURES_DIR,
    JOBS_FILE,
    os.path.abspath(__file__),
]

_DEFAULT_FIELDS = {
    "report_title": "Harness Test Report",
    "pqa_serial": "PQA-TEST-001",
    "gen_sn": "GEN-TEST-001",
    "site_address": "",
    "custom_text": "",
}


# ── helpers (mirror desktop/shell.py) ────────────────────────────────────────
def _fixture(name: str) -> str:
    """Resolve a fixture filename to an absolute path (absolute paths pass through)."""
    if os.path.isabs(name):
        return name
    return os.path.join(FIXTURES_DIR, name)


def _first_match(exts: tuple[str, ...]) -> str | None:
    if not os.path.isdir(FIXTURES_DIR):
        return None
    for name in sorted(os.listdir(FIXTURES_DIR)):
        if name.lower().endswith(exts) and not name.startswith("."):
            return os.path.join(FIXTURES_DIR, name)
    return None


def _load_frame(path: str):
    """Load CSV or WinScope XLS exactly as the shell's load_csv/load_winscope do."""
    import core.analysis as ca

    if path.lower().endswith((".xls", ".xlsx")):
        from desktop.xls_host import load_winscope_df
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return load_winscope_df(b64, os.path.basename(path))
    with open(path, "rb") as f:
        return ca.load_and_prepare_csv(io.BytesIO(f.read()))


def _build_config(overrides: dict):
    """Build AnalysisConfig the same way desktop/shell.py:_build_config does."""
    import dataclasses
    import core.analysis as ca

    cfg = ca.AnalysisConfig()
    valid = {f.name for f in dataclasses.fields(cfg)}
    for key, value in (overrides or {}).items():
        if key in valid and value is not None:
            setattr(cfg, key, value)
    return cfg


def _audit_docx_images(docx_path: str) -> dict:
    """Inspect image placement in the .docx without opening Word.

    Reports each inline image's rendered size (EMU→cm) and aspect, and flags
    stretched/oversized/near-zero pastes — a programmatic first pass before the
    visual check.
    """
    try:
        import docx  # python-docx
    except Exception as exc:  # noqa: BLE001
        return {"error": f"python-docx unavailable: {exc}"}

    EMU_PER_CM = 360000
    doc = docx.Document(docx_path)
    images = []
    for i, shape in enumerate(doc.inline_shapes):
        try:
            w_cm = round(shape.width / EMU_PER_CM, 2)
            h_cm = round(shape.height / EMU_PER_CM, 2)
        except Exception:  # noqa: BLE001
            w_cm = h_cm = None
        images.append({
            "index": i, "width_cm": w_cm, "height_cm": h_cm,
            "aspect": round(w_cm / h_cm, 2) if (w_cm and h_cm) else None,
        })
    # A4 text column tops out near 18cm; inject_images_to_word sizes pictures to
    # the page's content width, so only flag clearly-broken extremes: genuine
    # overflow (>19cm) or a collapsed/near-zero paste (<1cm).
    flags = []
    for im in images:
        w = im["width_cm"]
        if w is not None and w > 19:
            flags.append(f"image {im['index']} is {w}cm wide — overflows the page")
        if w is not None and w < 1:
            flags.append(f"image {im['index']} is {w}cm wide — suspiciously small")
    return {"count": len(images), "images": images, "flags": flags}


_SNAPSHOT_SLOT = __import__("re").compile(r"\{\{Snapshot_(\d+)\}\}")


def _count_template_snapshot_slots(template_path: str) -> int | None:
    """Count distinct ``{{Snapshot_N}}`` slots a Word template declares.

    The app renders one snapshot per detected event but can only place as many as
    the template has slots; surplus events are silently dropped. Counting the
    template's slots lets the harness flag that mismatch (events > slots) — and
    the opposite (empty slots, events < slots) — which the app does not warn on.
    """
    try:
        import docx
    except Exception:  # noqa: BLE001
        return None
    doc = docx.Document(template_path)
    seen: set[int] = set()

    def scan(paragraphs):
        for p in paragraphs:
            for m in _SNAPSHOT_SLOT.finditer("".join(r.text for r in p.runs)):
                seen.add(int(m.group(1)))

    scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return len(seen)


def _docx_to_pdf(docx_path: str, out_dir: str) -> str | None:
    """Convert .docx → PDF via LibreOffice for visual inspection.

    LibreOffice names its output after the input stem (``report.docx`` →
    ``report.pdf``), which would clobber the HTML-pipeline ``report.pdf``. So we
    convert into an isolated temp dir and then move the result to
    ``report.docx.pdf``.
    """
    if not _SOFFICE:
        return None
    import shutil
    import tempfile
    conv_dir = tempfile.mkdtemp(prefix="pqa_lo_", dir=out_dir)
    try:
        subprocess.run(
            [_SOFFICE, "--headless", "--convert-to", "pdf", "--outdir",
             conv_dir, docx_path],
            capture_output=True, text=True, timeout=120, check=False,
        )
        stem = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        produced = os.path.join(conv_dir, stem)
        if os.path.exists(produced):
            final = docx_path + ".pdf"
            os.replace(produced, final)
            return final
        return None
    except Exception:  # noqa: BLE001
        return None
    finally:
        shutil.rmtree(conv_dir, ignore_errors=True)


def _audit_docx_toc(docx_path: str) -> dict:
    """Report what the Word TOC will contain after Word refreshes it.

    LibreOffice doesn't refresh the TOC on headless convert, so the preview PDF's
    contents page is stale. Instead we read the doc's Heading-1..3 paragraphs (the
    exact set the TOC field includes) in document order and confirm updateFields is
    set — proving the TOC WILL be correct when opened in Word.
    """
    try:
        import docx
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
    except Exception as exc:  # noqa: BLE001
        return {"error": f"python-docx unavailable: {exc}"}

    doc = docx.Document(docx_path)
    headings = []
    for child in doc.element.body.iterchildren():
        if child.tag.split("}")[-1] != "p":
            continue
        p = Paragraph(child, doc)
        style = (p.style.name or "") if p.style is not None else ""
        if style.startswith("Heading") and p.text.strip():
            headings.append({"level": style, "text": p.text.strip()[:70]})
    update_fields = doc.settings.element.find(qn("w:updateFields")) is not None
    return {"update_fields_set": update_fields,
            "toc_headings": headings, "heading_count": len(headings)}


# ── one job ──────────────────────────────────────────────────────────────────
def run_job(job: dict) -> dict:
    """Run a single report job end-to-end; write artifacts; return its summary."""
    import core.analysis as ca
    from desktop.report_host import build_report

    name = job.get("name") or "report"
    csv_path = _fixture(job["csv"]) if job.get("csv") else _first_match((".csv", ".xls", ".xlsx"))
    template_path = _fixture(job["template"]) if job.get("template") else None
    t_start = job.get("time_start")
    t_end = job.get("time_end")
    fields = {**_DEFAULT_FIELDS, **(job.get("fields") or {})}
    config_overrides = job.get("config") or {}

    out_dir = os.path.join(OUTPUT_DIR, name)
    os.makedirs(out_dir, exist_ok=True)

    print(f"\n══ JOB: {name} ═══════════════════════════════════════")
    print(f"→ logger:   {csv_path}")
    print(f"→ template: {template_path or '(none — PDF/HTML only)'}")
    if t_start or t_end:
        print(f"→ window:   {t_start or '(open)'}  →  {t_end or '(open)'}")

    if not csv_path or not os.path.exists(csv_path):
        raise FileNotFoundError(f"[{name}] logger file not found: {csv_path}")
    if template_path and not os.path.exists(template_path):
        raise FileNotFoundError(f"[{name}] template not found: {template_path}")

    # 1. Load + validate (mirrors load_csv / load_winscope).
    df = _load_frame(csv_path)
    fmt = df.attrs.get("logger_format")

    # 2. Time-window slice + analyse (mirrors run_analysis).
    df_run = ca.filter_time_window(df, t_start, t_end)
    if df_run is None or df_run.empty:
        raise RuntimeError(f"[{name}] selected time window contains no data.")
    df_run.attrs.update(df.attrs)
    ok, errors, warnings = ca.validate_csv_format(df_run)
    print(f"→ format: {fmt} | rows in window: {len(df_run)}/{len(df)} | valid: {ok}")
    if errors:
        print(f"  validation errors: {errors}")

    cfg = _build_config(config_overrides)
    if fmt in ("miro", "winscope"):
        cfg.skip_interpolation = True
    df_proc, df_events = ca.perform_analysis(df_run, cfg)
    df_events = df_events.reset_index(drop=True)
    n_events = len(df_events)
    print(f"→ detected events: {n_events}")

    # Template snapshot-slot coverage: the app silently drops events with no slot
    # to land in, and leaves surplus slots blank — flag both so neither is missed.
    slot_flags: list[str] = []
    n_slots = _count_template_snapshot_slots(template_path) if template_path else None
    if n_slots is not None:
        print(f"→ template snapshot slots: {n_slots}")
        if n_events > n_slots:
            slot_flags.append(
                f"{n_events} events detected but template has only {n_slots} "
                f"snapshot slots — {n_events - n_slots} event(s) will be DROPPED "
                f"from the report.")
        elif n_events < n_slots:
            slot_flags.append(
                f"{n_events} events detected but template has {n_slots} snapshot "
                f"slots — {n_slots - n_events} slot(s) will render blank.")

    # 3. Build report artifacts (the real pipeline).
    want_docx = template_path is not None
    docx_b64 = None
    if want_docx:
        with open(template_path, "rb") as f:
            docx_b64 = base64.b64encode(f.read()).decode("ascii")
    include_ct = bool(job.get("include_compliance_table"))
    include_itic = bool(job.get("include_itic"))
    if include_ct or include_itic:
        print(f"→ toggles: compliance_table={include_ct} itic={include_itic}")
    params = {
        "fields": fields,
        "filename": "report",
        "outputs": {"pdf": True, "html": True, "docx": want_docx},
        "docx_template_b64": docx_b64,
        "include_compliance_table": include_ct,
        "include_itic": include_itic,
    }
    result = build_report(df_run, df_proc, df_events, cfg, params)
    artifacts = result.get("artifacts", {})

    # 4. Persist artifacts.
    written = []
    if "pdf_b64" in artifacts:
        p = os.path.join(out_dir, "report.pdf")
        with open(p, "wb") as f:
            f.write(base64.b64decode(artifacts["pdf_b64"]))
        written.append(p)
    if "html" in artifacts:
        p = os.path.join(out_dir, "report.html")
        with open(p, "w", encoding="utf-8") as f:
            f.write(artifacts["html"])
        written.append(p)
    docx_audit = None
    toc_audit = None
    if "docx_b64" in artifacts:
        p = os.path.join(out_dir, "report.docx")
        with open(p, "wb") as f:
            f.write(base64.b64decode(artifacts["docx_b64"]))
        written.append(p)
        docx_audit = _audit_docx_images(p)
        toc_audit = _audit_docx_toc(p)
        docx_pdf = _docx_to_pdf(p, out_dir)
        if docx_pdf:
            written.append(docx_pdf)

    # 5. Summary.
    summary = {
        "name": name,
        "inputs": {"logger": csv_path, "template": template_path,
                   "time_start": t_start, "time_end": t_end},
        "logger_format": fmt,
        "rows_in_window": int(len(df_run)),
        "rows_total": int(len(df)),
        "valid": bool(ok),
        "validation_errors": errors,
        "validation_warnings": warnings,
        "events_detected": int(n_events),
        "template_snapshot_slots": n_slots,
        "slot_coverage_flags": slot_flags,
        "report": {
            "filename": result.get("filename"),
            "pdf_ok": result.get("pdf_ok"),
            "warnings": result.get("warnings", []),
            "artifacts_present": sorted(artifacts.keys()),
        },
        "docx_image_audit": docx_audit,
        "docx_toc_audit": toc_audit,
        "outputs_written": written,
    }
    with open(os.path.join(out_dir, "summary.json"), "w") as f:
        json.dump(summary, f, indent=2, default=str)

    print(f"── {name}: pdf_ok={result.get('pdf_ok')} events={n_events} ──")
    for w in result.get("warnings", []):
        print(f"  • {w}")
    for fl in slot_flags:
        print(f"  ⚠ {fl}")
    if docx_audit and docx_audit.get("flags"):
        for fl in docx_audit["flags"]:
            print(f"  ⚠ {fl}")
    if toc_audit and not toc_audit.get("error"):
        print(f"  TOC: {toc_audit['heading_count']} headings, "
              f"updateFields={toc_audit['update_fields_set']}")
        for h in toc_audit["toc_headings"]:
            print(f"      · {h['text']}")
    for w in written:
        print(f"  {w}")
    return summary


# ── job discovery ─────────────────────────────────────────────────────────────
def _load_jobs(args) -> list[dict]:
    """Build the job list from CLI flags, else from jobs.json, else a default."""
    if args.csv or args.template or args.time_start or args.time_end or args.name:
        return [{
            "name": args.name or "report",
            "csv": args.csv,
            "template": args.template,
            "time_start": args.time_start,
            "time_end": args.time_end,
        }]
    if os.path.exists(JOBS_FILE):
        with open(JOBS_FILE) as f:
            jobs = (json.load(f) or {}).get("jobs", [])
        if args.job:
            jobs = [j for j in jobs if j.get("name") == args.job]
            if not jobs:
                raise SystemExit(f"no job named '{args.job}' in {JOBS_FILE}")
        return jobs
    # No jobs.json, no flags → single default run over the first fixtures found.
    return [{"name": "report"}]


def _run_all(jobs: list[dict]) -> int:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    rc = 0
    for job in jobs:
        try:
            run_job(job)
        except Exception as exc:  # noqa: BLE001 — keep going so other jobs still run
            rc = 1
            print(f"  ✗ job '{job.get('name')}' failed: {exc}")
    return rc


# ── watch mode ─────────────────────────────────────────────────────────────────
def _snapshot_mtimes() -> dict:
    """Latest mtime per watched path (recursing into dirs)."""
    snap = {}
    for root in _WATCH_SOURCES:
        if os.path.isfile(root):
            snap[root] = os.path.getmtime(root)
        elif os.path.isdir(root):
            for dirpath, _dirs, files in os.walk(root):
                if "__pycache__" in dirpath:
                    continue
                for fn in files:
                    if fn.endswith(".pyc") or fn.startswith("."):
                        continue
                    fp = os.path.join(dirpath, fn)
                    try:
                        snap[fp] = os.path.getmtime(fp)
                    except OSError:
                        pass
    return snap


def _watch(args, interval: float = 1.0) -> int:
    print("⟳ watch mode — re-runs on changes to fixtures/, jobs.json, or the "
          "report pipeline. Ctrl-C to stop.")
    _run_all(_load_jobs(args))
    last = _snapshot_mtimes()
    try:
        while True:
            time.sleep(interval)
            now = _snapshot_mtimes()
            if now != last:
                changed = [p for p in now if now.get(p) != last.get(p)] + \
                          [p for p in last if p not in now]
                print(f"\n⟳ change detected ({len(changed)} file(s)) — re-running")
                for c in changed[:5]:
                    print(f"    ~ {c.replace(_REPO_ROOT + '/', '')}")
                _run_all(_load_jobs(args))  # reload jobs.json each cycle
                last = _snapshot_mtimes()
    except KeyboardInterrupt:
        print("\n⟳ watch stopped")
        return 0


def main() -> int:
    # Line-buffer stdout so progress is visible live even when redirected to a
    # file or pipe (Python block-buffers a non-TTY by default) — important for
    # --watch, where the user is tailing the log.
    try:
        sys.stdout.reconfigure(line_buffering=True)
    except Exception:  # noqa: BLE001
        pass

    ap = argparse.ArgumentParser(description="Headless PQA report test harness")
    ap.add_argument("--job", help="run only the named job from jobs.json")
    ap.add_argument("--watch", action="store_true", help="re-run on file changes")
    # single-shot CLI mode (bypasses jobs.json)
    ap.add_argument("--csv", help="logger CSV/XLS (default: first in fixtures/)")
    ap.add_argument("--template", help="Word .docx template")
    ap.add_argument("--time-start", dest="time_start", help="ISO window start")
    ap.add_argument("--time-end", dest="time_end", help="ISO window end")
    ap.add_argument("--name", help="output subfolder name for single-shot mode")
    args = ap.parse_args()

    if args.watch:
        return _watch(args)
    return _run_all(_load_jobs(args))


if __name__ == "__main__":
    raise SystemExit(main())
