"""Persistent Word-template library (desktop only).

The Streamlit app let users upload several ``.docx`` report templates, kept them
in ``uploads/templates/``, listed them with a remove button, and picked one from
a dropdown. The desktop app reproduces that workflow here, but stores templates
in the durable per-user app-data dir (``%APPDATA%\\PQA\\templates`` on Windows,
``~/Library/Application Support/PQA/templates`` on macOS, ``~/.local/share/PQA``
on Linux — see :func:`desktop.usage_log.data_dir`) so they survive app restarts
and reinstalls, unlike the ephemeral Streamlit ``uploads/`` dir.

Each template is scanned for its ``{{Snapshot_N}}`` placeholders so the frontend
can run the same pre-flight completeness check the Streamlit app did (warn when a
template has fewer snapshot slots than the analysis detected events). The scan
needs python-docx, which is imported lazily so this module stays importable under
the lean CI pytest job (pandas/numpy/pytest only).
"""
from __future__ import annotations

import base64
import os
import re

from desktop import usage_log

# Matches the snapshot placeholder tokens the report templates use. The index is
# captured so the caller can compare a template's slots against the event count.
_SNAPSHOT_RE = re.compile(r"\{\{Snapshot_(\d+)\}\}")

_FORBIDDEN = '<>:"/\\|?*'


def _safe_name(name: str) -> str:
    """Filesystem-safe basename: strip separators/control chars, keep ``.docx``.

    Mirrors ``report_host._safe_name`` but always preserves the extension so the
    stored file stays a valid Word document the picker can re-open.
    """
    name = os.path.basename(str(name or "")).strip()
    stem = os.path.splitext(name)[0]  # drop any original extension; .docx is forced
    cleaned = "".join("_" if ch in _FORBIDDEN or ord(ch) < 32 else ch for ch in stem)
    cleaned = cleaned.strip(" .")
    return (cleaned or "template") + ".docx"


def templates_dir() -> str:
    """Path to the persistent template library, created on first use."""
    path = os.path.join(usage_log.data_dir(), "templates")
    os.makedirs(path, exist_ok=True)
    return path


def resolve(name: str) -> str | None:
    """Resolve a stored template name to an absolute path inside the library.

    Returns ``None`` when the name is empty, escapes the library dir, or the file
    does not exist — so a stale/forged name can never read an arbitrary file.
    """
    if not name:
        return None
    base = templates_dir()
    path = os.path.normpath(os.path.join(base, _safe_name(name)))
    if os.path.commonpath([base, path]) != base:
        return None
    return path if os.path.isfile(path) else None


def scan_snapshot_indices(path: str) -> list[int]:
    """Sorted list of ``N`` for every ``{{Snapshot_N}}`` placeholder in a .docx.

    Scans paragraph and table-cell text. Returns ``[]`` on any read error so a
    corrupt upload degrades to "no slots" rather than raising — the caller treats
    that as a (loud) completeness shortfall, not a crash.
    """
    try:
        from docx import Document  # lazy: python-docx only needed to scan
    except Exception:  # noqa: BLE001 — CI without python-docx
        return []

    found: set[int] = set()
    try:
        doc = Document(path)

        def _collect(text: str) -> None:
            for m in _SNAPSHOT_RE.findall(text or ""):
                found.add(int(m))

        for para in doc.paragraphs:
            _collect(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _collect(cell.text)
    except Exception:  # noqa: BLE001 — unreadable docx → no slots
        return []
    return sorted(found)


def _info(path: str) -> dict:
    """Metadata dict for one stored template (name, size, snapshot slots)."""
    idxs = scan_snapshot_indices(path)
    return {
        "name": os.path.basename(path),
        "size": os.path.getsize(path),
        "snapshot_indices": idxs,
        "snapshot_max": max(idxs) if idxs else 0,
    }


def list_templates() -> list[dict]:
    """All stored templates as metadata dicts, sorted by name."""
    base = templates_dir()
    out = []
    for fname in sorted(os.listdir(base)):
        if fname.lower().endswith(".docx"):
            out.append(_info(os.path.join(base, fname)))
    return out


def save_template(filename: str, b64: str) -> list[dict]:
    """Persist a base64 ``.docx`` upload, returning the updated library list.

    The name is sanitised; an existing template with the same name is overwritten
    (re-uploading replaces, matching the Streamlit behaviour).
    """
    if not b64:
        raise ValueError("save_template requires b64 content")
    path = os.path.join(templates_dir(), _safe_name(filename))
    with open(path, "wb") as f:
        f.write(base64.b64decode(b64))
    return list_templates()


def delete_template(name: str) -> list[dict]:
    """Remove a stored template by name, returning the updated library list."""
    path = resolve(name)
    if path and os.path.isfile(path):
        os.remove(path)
    return list_templates()
