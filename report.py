"""
Report generation for Power Quality Analysis.

Handles Word template injection, placeholder mapping, and PDF conversion.
"""

import os
import glob
import io
import logging
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

log = logging.getLogger(__name__)


_PDF_SCRIPT = r"""
import sys, os, io, tempfile
docx_path, pdf_path = sys.argv[1], sys.argv[2]

def try_libreoffice():
    import shutil as _shutil
    import subprocess as _sp

    candidates = [
        'libreoffice',
        'soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/bin/libreoffice',
        '/usr/bin/soffice',
        '/snap/bin/libreoffice',
    ]
    lo_bin = None
    for c in candidates:
        try:
            r = _sp.run([c, '--version'], capture_output=True, timeout=5)
            if r.returncode == 0:
                lo_bin = c
                break
        except Exception:
            continue
    if not lo_bin:
        raise RuntimeError('LibreOffice not found')

    out_dir = os.path.dirname(os.path.abspath(pdf_path))
    _sp.run(
        [lo_bin, '--headless', '--convert-to', 'pdf', '--outdir', out_dir,
         os.path.abspath(docx_path)],
        check=True, capture_output=True, timeout=90,
    )
    # LibreOffice names the output after the input basename
    expected = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
    if os.path.exists(expected) and os.path.abspath(expected) != os.path.abspath(pdf_path):
        _shutil.move(expected, pdf_path)

def try_docx2pdf():
    from docx2pdf import convert
    convert(docx_path, pdf_path)
    if not os.path.exists(pdf_path):
        raise RuntimeError('docx2pdf ran but no output file produced')

def try_weasyprint():
    import mammoth, weasyprint
    with open(docx_path, 'rb') as f:
        result = mammoth.convert_to_html(f)
    html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        '<style>'
        'body{font-family:Arial,sans-serif;font-size:11pt;margin:2cm;line-height:1.4}'
        'img{max-width:100%;height:auto;display:block;margin:10pt auto}'
        'table{width:100%;border-collapse:collapse;margin:10pt 0}'
        'td,th{border:1px solid #ccc;padding:4pt 6pt}'
        'p{margin:4pt 0}'
        '</style></head><body>' + result.value + '</body></html>'
    )
    weasyprint.HTML(string=html).write_pdf(pdf_path)

def try_fpdf2():
    from docx import Document
    from docx.oxml.ns import qn
    from fpdf import FPDF
    from PIL import Image

    doc = Document(docx_path)
    pdf = FPDF(format='A4')
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Extract embedded images from the docx zip
    import zipfile
    img_map = {}
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                ext = os.path.splitext(name)[1].lower()
                if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp'):
                    img_map[name] = z.read(name)

    # Map relationship IDs to image data
    rel_to_img = {}
    with zipfile.ZipFile(docx_path) as z:
        if 'word/_rels/document.xml.rels' in z.namelist():
            import xml.etree.ElementTree as ET
            rels_xml = z.read('word/_rels/document.xml.rels')
            root = ET.fromstring(rels_xml)
            for rel in root:
                target = rel.get('Target', '')
                if 'media/' in target:
                    rid = rel.get('Id')
                    img_key = 'word/' + target.lstrip('/')
                    if img_key in img_map:
                        rel_to_img[rid] = img_map[img_key]

    tmp_imgs = []
    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    for para in doc.paragraphs:
        # Check for inline images
        for elem in para._element.iter():
            if elem.tag.endswith('}blip'):
                rid = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rid and rid in rel_to_img:
                    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    tmp.write(rel_to_img[rid])
                    tmp.close()
                    tmp_imgs.append(tmp.name)
                    try:
                        with Image.open(tmp.name) as im:
                            w, h = im.size
                        ratio = h / w if w > 0 else 1
                        img_w = min(page_w, page_w)
                        img_h = img_w * ratio
                        if pdf.get_y() + img_h > pdf.page_break_trigger:
                            pdf.add_page()
                        pdf.image(tmp.name, x=pdf.l_margin, w=img_w)
                        pdf.ln(3)
                    except Exception:
                        pass

        text = para.text.strip()
        if not text:
            pdf.ln(3)
            continue

        style = para.style.name if para.style else ''
        if 'Heading 1' in style:
            pdf.set_font('Helvetica', 'B', 16)
        elif 'Heading 2' in style:
            pdf.set_font('Helvetica', 'B', 13)
        elif 'Heading' in style:
            pdf.set_font('Helvetica', 'B', 11)
        else:
            pdf.set_font('Helvetica', '', 11)

        pdf.multi_cell(0, 6, text)
        pdf.ln(1)

    pdf.output(pdf_path)

    for t in tmp_imgs:
        try: os.unlink(t)
        except: pass

# Try converters in order
for converter in [try_libreoffice, try_docx2pdf, try_weasyprint, try_fpdf2]:
    try:
        converter()
        if os.path.exists(pdf_path):
            print(f'PDF conversion succeeded via {converter.__name__}')
            sys.exit(0)
    except Exception as e:
        print(f'{converter.__name__} failed: {e}', file=sys.stderr)

sys.exit(1)
"""


def convert_to_pdf(docx_path, pdf_path, timeout=45):
    """
    Run PDF conversion in a subprocess with a timeout.
    Returns (success: bool, log: str) — log contains converter diagnostics
    whether conversion succeeded or failed.
    """
    import sys
    import subprocess

    try:
        result = subprocess.run(
            [sys.executable, "-c", _PDF_SCRIPT,
             os.path.abspath(docx_path), os.path.abspath(pdf_path)],
            timeout=timeout,
            capture_output=True,
            text=True,
        )
        success = os.path.exists(pdf_path)
        log_lines = []
        if result.stdout.strip():
            log_lines.append(result.stdout.strip())
        if result.stderr.strip():
            log_lines.append(result.stderr.strip())
        if not log_lines:
            log_lines.append("(no output from converter)")
        log = "\n".join(log_lines)
        return success, log
    except subprocess.TimeoutExpired:
        return False, f"Timed out after {timeout}s — converter hung."
    except Exception as e:
        return False, f"Subprocess error: {e}"


def get_placeholder_map(client_name, config_values, df=None,
                        graph_dir="output/Graphs",
                        snapshot_dir="output/Snapshots",
                        image_dir="output/Images"):
    """
    Build a mapping of {{placeholder}} -> file path or text value.

    Parameters:
        client_name: string identifier for the client/report
        config_values: dict with keys like 'report_title', 'gen_sn', 'site_address', 'custom_text'
        df: optional DataFrame with Timestamp column for date/time extraction
        graph_dir, snapshot_dir, image_dir: output directories
    """
    placeholder_map = {}

    # Compliance table
    table_path = os.path.join(image_dir, f"{client_name}_table.png")
    if os.path.exists(table_path):
        placeholder_map["{{Compliance_Table}}"] = table_path

    # ITIC (CBEMA) curve — only emitted when the toggle rendered the image.
    itic_path = os.path.join(graph_dir, f"{client_name}_ITIC_Curve.jpeg")
    if os.path.exists(itic_path):
        placeholder_map["{{ITIC_Curve}}"] = itic_path

    # Metric graphs
    metrics = {
        "Avg_Voltage_LL": "{{Avg_Voltage_LL}}",
        "Avg_kW": "{{Avg_kW}}",
        "Avg_Current": "{{Avg_Current}}",
        "Avg_Frequency": "{{Avg_Frequency}}",
        "Avg_THD_F": "{{Avg_THD_F}}",
        "Avg_PF": "{{Avg_PF}}",
    }
    for metric, placeholder in metrics.items():
        graph_path = os.path.join(graph_dir, f"{client_name}_{metric}.jpeg")
        if os.path.exists(graph_path):
            placeholder_map[placeholder] = graph_path

    # Snapshots
    snapshot_files = sorted(glob.glob(os.path.join(snapshot_dir, f"snap_{client_name}_*.jpeg")))
    for i, file_path in enumerate(snapshot_files, start=1):
        placeholder_map[f"{{{{Snapshot_{i}}}}}"] = file_path

    # Text fields
    placeholder_map["{{Report_Title}}"] = config_values.get("report_title", "")
    placeholder_map["{{PQID}}"] = config_values.get("pqa_serial", "")
    placeholder_map["{{Gen_SN}}"] = config_values.get("gen_sn", "")
    placeholder_map["{{Site_Address}}"] = config_values.get("site_address", "")
    placeholder_map["{{Custom_Field}}"] = config_values.get("custom_text", "")

    # Date/time extracted from CSV data (not today's date)
    if df is not None and not df.empty and "Timestamp" in df.columns:
        min_ts = df["Timestamp"].min()
        max_ts = df["Timestamp"].max()
        fmt_dt = "%d/%m/%Y %I:%M:%S %p"
        placeholder_map["{{Start Time}}"] = min_ts.strftime(fmt_dt)
        placeholder_map["{{End Time}}"] = max_ts.strftime(fmt_dt)

        # Extract {{Date}} directly from the CSV "Date" column when available.
        # This avoids inheriting wrong dates from a "PC Time" column that may
        # have a misconfigured clock. Fallback is Timestamp.min() (also CSV-
        # sourced). No fallback to the computer's current date in either path.
        date_str = None
        if "Date" in df.columns:
            raw_date = df["Date"].dropna()
            if not raw_date.empty:
                try:
                    parsed = pd.to_datetime(raw_date.iloc[0], dayfirst=True)
                    date_str = parsed.strftime("%d/%m/%Y")
                except Exception:
                    pass
        if date_str is None:
            date_str = min_ts.strftime("%d/%m/%Y")
        placeholder_map["{{Date}}"] = date_str

    return placeholder_map


# Image placeholders that get_placeholder_map only emits when the underlying
# file exists. When a template has more snapshot/metric slots than the analysis
# produced (e.g. a 6-step template run against a 5-event test), the surplus
# placeholders are never in the map and would otherwise survive as literal
# {{Snapshot_6}} text in the rendered .docx. We blank these leftovers so a Word
# report never shows raw template syntax — mirrors report_host's HTML strip.
_UNUSED_IMAGE_PLACEHOLDER = re.compile(
    r"\{\{(?:Snapshot_\d+|Avg_[A-Za-z_]+|Compliance_Table|ITIC_Curve)\}\}")


def _element_all_text(elem):
    """Concatenate every ``w:t`` text node inside a body element (p or tbl)."""
    return "".join(t.text or "" for t in elem.findall(".//" + qn("w:t")))


def _is_section_heading(p_elem, doc):
    """True when a paragraph reads as a section heading.

    Checks the style name first ('Heading N' / 'Title'), then the **outline
    level** — set either directly on the paragraph or inherited from its style —
    so a heading the TOC picks up via outline level still counts even when its
    style isn't literally named 'Heading N'. (Some templates author the first
    'Load Step 1' heading in a plain style, which is why the style name alone is
    not enough.)
    """
    from docx.text.paragraph import Paragraph

    p = Paragraph(p_elem, doc)
    try:
        name = (p.style.name if p.style is not None else "") or ""
    except Exception:  # noqa: BLE001
        name = ""
    if name.startswith("Heading") or name == "Title":
        return True
    candidates = [p_elem.find(qn("w:pPr"))]
    try:
        candidates.append(p.style.element.find(qn("w:pPr")))
    except Exception:  # noqa: BLE001
        pass
    for pPr in candidates:
        if pPr is None:
            continue
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            try:
                return int(ol.get(qn("w:val"))) < 9
            except (TypeError, ValueError):
                pass
    return False


def _set_page_break_before(elem):
    """Make ``elem`` start on a new page (idempotent).

    A paragraph gets a ``w:pageBreakBefore`` property. Tables can't carry that
    property, so a page-break spacer paragraph is inserted immediately before the
    table instead.
    """
    tag = elem.tag.split("}")[-1]
    if tag == "p":
        pPr = elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            elem.insert(0, pPr)
        if pPr.find(qn("w:pageBreakBefore")) is None:
            pPr.append(OxmlElement("w:pageBreakBefore"))
    elif tag == "tbl":
        brk = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pPr.append(OxmlElement("w:pageBreakBefore"))
        brk.append(pPr)
        elem.addprevious(brk)


def _find_snapshot_anchor(doc):
    """Body element that new 'after results / before snapshots' content goes before.

    Finds the first ``{{Snapshot_*}}`` block, then walks back to the heading that
    introduces it (e.g. 'Load Step 1') so injected sections land *above* that
    heading. The first snapshot heading is sometimes authored in a plain
    (non-'Heading') style, so we treat the topmost paragraph of the contiguous
    non-blank run immediately preceding the snapshot as the heading, and stop
    early on a real heading (by :func:`_is_section_heading`). Returns that
    paragraph; failing that the snapshot block itself; ``None`` when the template
    has no snapshots (caller appends at the end). Must be called BEFORE
    placeholders are replaced, while the ``{{Snapshot_1}}`` marker text still
    exists.
    """
    snap_re = re.compile(r"\{\{Snapshot_\d+\}\}")
    children = list(doc.element.body.iterchildren())
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] not in ("p", "tbl"):
            continue
        if not snap_re.search(_element_all_text(child)):
            continue
        anchor = None
        j = i - 1
        while j >= 0:
            prev = children[j]
            if prev.tag.split("}")[-1] != "p":
                break  # a table or other block — boundary above the heading group
            if not _element_all_text(prev).strip():
                break  # blank spacer line — boundary above the heading group
            anchor = prev                       # a non-blank paragraph in the group
            if _is_section_heading(prev, doc):
                break                            # a real heading — section start
            j -= 1
        return anchor if anchor is not None else child
    return None


# Time-series metric plots live behind these placeholders ({{Avg_Voltage_LL}},
# {{Avg_kW}}, …). The Compliance Table / ITIC Curve are injected directly after
# the LAST of them — i.e. right after the time-series plots block.
_TIMESERIES_PLOT_RE = re.compile(r"\{\{Avg_[A-Za-z0-9_]+\}\}")


def _find_results_end_anchor(doc):
    """Body block to insert the ITIC/compliance sections *before* — the first
    block after the time-series plots.

    Anchoring off the plots (the well-defined ``{{Avg_*}}`` placeholders) rather
    than off the snapshot headings makes placement template-independent: the
    sections always land directly after the time-series plots, on their own page,
    no matter how the snapshot section that follows is styled. Locates the LAST
    ``{{Avg_*}}`` placeholder and returns the next meaningful sibling (skipping
    blank spacer paragraphs). Falls back to the snapshot-heading anchor when a
    template carries no plot placeholders; ``None`` means append at the end.
    Must run BEFORE placeholders are replaced, while the markers still exist.
    """
    children = list(doc.element.body.iterchildren())
    last_plot = -1
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] not in ("p", "tbl"):
            continue
        if _TIMESERIES_PLOT_RE.search(_element_all_text(child)):
            last_plot = i
    if last_plot < 0:
        return _find_snapshot_anchor(doc)  # no plots — fall back to snapshot anchor
    j = last_plot + 1
    while j < len(children):
        nxt = children[j]
        if nxt.tag.split("}")[-1] == "p" and not _element_all_text(nxt).strip():
            j += 1
            continue  # skip blank spacer paragraphs trailing the last plot
        return nxt
    return None  # the plots are the last content — append at the end


def _insert_image_sections(doc, anchor, sections, content_width):
    """Insert ``(heading, image_path)`` sections before ``anchor`` (append if None).

    Each section is a Heading-1 paragraph (so the TOC field lists it) followed by
    a left-aligned full-width picture. The whole injected block starts on a fresh
    page (page-break before the first section only), and the remaining sections
    flow on after it — so e.g. the Compliance Table and ITIC Curve share one page
    rather than each taking its own, while still landing above the first snapshot
    heading. Returns the list of heading titles actually inserted (those whose
    image existed).
    """
    added = []
    for title, image_path in sections:
        if not image_path or not os.path.exists(image_path):
            continue
        # Build at the end of the doc, then relocate before the anchor. Inserting
        # heading-then-image in order keeps them adjacent and correctly ordered.
        heading = doc.add_heading(title, level=1)
        # Only the first inserted section breaks to a new page; later ones flow
        # on the same page so the injected sections group together.
        heading.paragraph_format.page_break_before = (len(added) == 0)
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        img_para.add_run().add_picture(image_path, width=content_width)
        if anchor is not None:
            anchor.addprevious(heading._p)
            anchor.addprevious(img_para._p)
        added.append(title)
    # Resume the snapshots cleanly on a new page after the injected block, so the
    # first snapshot is moved to the next page intact instead of being cramped
    # onto the tail of the Compliance/ITIC page and clipped.
    if added and anchor is not None:
        _set_page_break_before(anchor)
    return added


def enable_update_fields(doc):
    """Flag all fields (incl. the TOC) to refresh when the document is next opened.

    Word honours this on open — refreshing the TOC's page numbers after we inject
    new headed sections. (LibreOffice ignores it on headless convert; the contents
    page there stays stale, but the body content is correct.)
    """
    settings = doc.settings.element
    upd = settings.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        settings.insert(0, upd)
    upd.set(qn("w:val"), "true")


def inject_images_to_word(template_stream, placeholder_map, *,
                          extra_sections=None, update_fields=False):
    """
    Replace {{placeholders}} in a Word document with images or text.

    Parameters:
        template_stream: file path string or file-like BytesIO object
        placeholder_map: dict of {{key}} -> file path or text value
        extra_sections: optional list of (heading_text, placeholder_key, image_path)
            to inject as new Heading-1 sections after the results/time-series block
            and before the per-event snapshots (used for the Compliance Table / ITIC
            Curve toggles). A section is injected only when the template does NOT
            already contain placeholder_key (otherwise normal replacement fills it
            in place).
        update_fields: when True, mark the document's fields (incl. the TOC) to
            refresh on next open, so the contents page picks up injected sections.

    Returns:
        Document object with replacements applied
    """
    doc = Document(template_stream)

    # Derive usable content width from the first section's page geometry.
    # A small safety buffer (0.15 in) is subtracted so images never reach the
    # exact margin edge — LibreOffice's PDF renderer can introduce sub-mm
    # offsets that clip a full-width image on the right side.
    try:
        _sec = doc.sections[0]
        _content_width = _sec.page_width - _sec.left_margin - _sec.right_margin - Inches(0.15)
    except Exception:
        _content_width = Inches(6.35)  # safe fallback

    # Inject sections (Compliance Table / ITIC) BEFORE replacing placeholders, so
    # the {{Snapshot_1}} anchor text is still present to locate the
    # results→snapshots boundary. Each requested section is (title, placeholder_key,
    # image_path): if the template already contains that placeholder we skip the
    # section (normal replacement will fill it in place); otherwise we inject.
    if extra_sections:
        body_text = _element_all_text(doc.element.body)
        to_inject = [(title, path) for (title, ph_key, path) in extra_sections
                     if path and os.path.exists(path) and ph_key not in body_text]
        if to_inject:
            # Anchor off the END of the time-series plots (template-independent),
            # not off the snapshot heading. _insert_image_sections page-breaks the
            # block onto its own page and page-breaks the following content too, so
            # the sections sit cleanly after the plots without disturbing the rest.
            anchor = _find_results_end_anchor(doc)
            _insert_image_sections(doc, anchor, to_inject, _content_width)

    def apply_strict_formatting(paragraph):
        p_format = paragraph.paragraph_format
        p_format.line_spacing = 1.0
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.keep_with_next = False
        p_format.left_indent = Inches(0)
        p_format.right_indent = Inches(0)
        p_format.first_line_indent = Inches(0)
        # Also clear the indent at the XML level so style-inherited indents
        # don't survive into LibreOffice's PDF renderer.
        pPr = paragraph._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)
        ind_elem = OxmlElement("w:ind")
        ind_elem.set(qn("w:left"), "0")
        ind_elem.set(qn("w:right"), "0")
        ind_elem.set(qn("w:firstLine"), "0")
        pPr.append(ind_elem)

    def _collapse_trailing_empties(para_elem, max_keep=1):
        """
        Remove consecutive empty paragraphs that immediately follow para_elem,
        keeping at most max_keep of them.  Templates often use long runs of empty
        paragraphs for spacing; after an image is injected those lines overflow
        onto the next page and produce a blank page.
        """
        body = para_elem.getparent()
        if body is None:
            return
        children = list(body)
        try:
            start = children.index(para_elem) + 1
        except ValueError:
            return
        empties = []
        for child in children[start:]:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "p":
                break
            text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
            if text:
                break
            empties.append(child)
        for elem in empties[max_keep:]:
            body.remove(elem)

    def preserve_paragraph_spacing(paragraph):
        """Ensure paragraphs have explicit spacing for LibreOffice PDF conversion."""
        p_format = paragraph.paragraph_format
        # Only set if not already set (preserve template intent)
        if p_format.space_before is None or p_format.space_before == Pt(0):
            p_format.space_before = Pt(3)
        if p_format.space_after is None or p_format.space_after == Pt(0):
            p_format.space_after = Pt(3)

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            # Build full paragraph text from all runs (handles split placeholders)
            full_text = "".join([run.text for run in paragraph.runs])
            new_text = full_text
            has_image = False
            image_path = None

            for key, value in placeholder_map.items():
                if key not in new_text:
                    continue

                if isinstance(value, str) and value.lower().endswith((".png", ".jpg", ".jpeg")) and os.path.exists(value):
                    # Image replacement: mark for processing after text replacements
                    has_image = True
                    image_path = value
                else:
                    # Text replacement: apply to new_text (which accumulates all replacements)
                    new_text = new_text.replace(key, str(value))

            # Blank any image/metric placeholders the map never filled (template
            # has more slots than this analysis produced) so no literal
            # {{Snapshot_N}} survives into the rendered document.
            if not has_image:
                new_text = _UNUSED_IMAGE_PLACEHOLDER.sub("", new_text)

            # Now apply the accumulated changes to the paragraph
            if has_image and image_path:
                # Image replacement: clear paragraph and add image
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                apply_strict_formatting(paragraph)
                run = paragraph.add_run()
                run.add_picture(image_path, width=_content_width)
                _collapse_trailing_empties(paragraph._p)
            elif new_text != full_text:
                # Text replacements were made: clear runs and rebuild once
                for run in list(paragraph.runs):
                    run._element.getparent().remove(run._element)

                # Add new run with replaced text, preserving font
                run = paragraph.add_run(new_text)
                run.font.name = "Arial"
                run.font.size = Pt(11)  # Use 11pt for body text, not 12pt

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    if update_fields:
        enable_update_fields(doc)

    return doc


def generate_docx(template_path, placeholder_map, output_name="PQA_Report"):
    """
    Inject placeholders into a Word template and save the .docx.
    Returns the path to the saved .docx file.
    """
    docx_path = f"{output_name}.docx"
    with open(template_path, "rb") as f:
        doc = inject_images_to_word(io.BytesIO(f.read()), placeholder_map)
    doc.save(docx_path)
    return docx_path


def generate_report(template_path, placeholder_map, output_name="PQA_Report"):
    """
    Full pipeline: inject placeholders into Word template and optionally convert to PDF.
    Returns dict with 'docx' and optionally 'pdf' keys pointing to file paths.
    """
    docx_path = generate_docx(template_path, placeholder_map, output_name)
    result = {"docx": docx_path}
    pdf_path = f"{output_name}.pdf"
    if convert_to_pdf(docx_path, pdf_path):
        result["pdf"] = pdf_path
    return result
